from docx import Document
from django.http import HttpResponse
from io import BytesIO
from django.utils import timezone
from django.utils.timezone import now
import locale
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import os
from django.conf import settings
from django.db.models import Prefetch
import logging
from .models import Crr, Veiculo, Arrendatario
from notificacao.models import NumeroEdital
from django.db import transaction



logger = logging.getLogger(__name__)

TEMPLATE_DOCX_PATH = os.path.join(settings.MEDIA_ROOT, 'modelo_edital.docx')

def obter_proximo_numero_edital():
    """Obtém e incrementa o número do edital de forma atômica"""
    try:
        with transaction.atomic():
            obj = NumeroEdital.objects.select_for_update().get(id=1)
            obj.numero += 1
            obj.save()
            return obj
    except NumeroEdital.DoesNotExist:
        obj = NumeroEdital.objects.create(id=1, numero=1)
        return obj

def gerar_edital_docx(crrs_ids):
    """
    Gera um documento DOCX para uma lista de IDs de CRRs
    """
    # Otimiza as queries com prefetch_related e select_related
    crrs = Crr.objects.filter(id__in=crrs_ids).select_related(
        'notificacao'
    ).prefetch_related(
        Prefetch('veiculo', queryset=Veiculo.objects.all()),
        Prefetch('arrendatarios', queryset=Arrendatario.objects.select_related('arrendatario'))
    )

    if not crrs.exists():
        raise ValueError("Nenhum CRR válido foi selecionado")

    try:
        doc = Document(TEMPLATE_DOCX_PATH)
    except Exception as e:
        logger.error(f"Erro ao carregar template DOCX: {e}")
        raise FileNotFoundError(f"Template DOCX não encontrado em {TEMPLATE_DOCX_PATH}") from e

    # Configura locale para data em português
    try:
        locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')
    except locale.Error:
        try:
            locale.setlocale(locale.LC_TIME, 'pt_BR')
        except locale.Error:
            pass

    # Obtém número do edital
    numero_edital_obj = obter_proximo_numero_edital()
    numero_edital = str(numero_edital_obj.numero - 1)  # Usa o número antes do incremento
    data_formatada = now().strftime('%d DE %B DE %Y').upper()

    # Substitui placeholders no documento
    for paragraph in doc.paragraphs:
        original_text = "".join([run.text for run in paragraph.runs])
        replaced_text = original_text.replace("{{DATA_ATUAL}}", data_formatada).replace("{{NUMERO_EDITAL}}", numero_edital)
        
        if original_text != replaced_text:
            if paragraph.runs:
                first_run = paragraph.runs[0]
                font = {
                    'name': first_run.font.name,
                    'size': first_run.font.size,
                    'bold': first_run.bold,
                    'italic': first_run.italic,
                    'underline': first_run.underline
                }
                
                # Limpa runs existentes
                while paragraph.runs:
                    paragraph._element.remove(paragraph.runs[-1]._element)
                
                # Adiciona novo run com formatação
                new_run = paragraph.add_run(replaced_text)
                new_run.font.name = font['name']
                new_run.font.size = font['size']
                new_run.bold = font['bold']
                new_run.italic = font['italic']
                new_run.underline = font['underline']

                try:
                    if new_run._element.rPr is None:
                        new_run._element.add_rPr()
                    if new_run._element.rPr.rFonts is None:
                        new_run._element.rPr.add_rFonts()
                    new_run._element.rPr.rFonts.set(qn('w:eastAsia'), font['name'] or 'Verdana')
                except Exception:
                    pass

            if "{{DATA_ATUAL}}" in original_text or "{{NUMERO_EDITAL}}" in original_text:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Adiciona tabela com dados dos CRRs
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Cabeçalho
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'PLACA/CHASSI'
    hdr_cells[1].text = 'MARCA/MODELO'
    hdr_cells[2].text = 'RESPONSÁVEL'
    hdr_cells[3].text = 'AGENTE FINACEIRO/ARRENDATÁRIO'

    # Dados
    for crr in crrs:
        row_cells = table.add_row().cells
        
        # Veículo
        veiculo = crr.veiculo.first() if crr.veiculo.exists() else None
        placa = veiculo.placa if veiculo else ''
        marca = veiculo.marca if veiculo else ''
        modelo = veiculo.modelo if veiculo else ''
        
        # Notificação
        notificacao = getattr(crr, 'notificacao', None)
        responsavel = notificacao.destinatario if notificacao else ''
        
        # Arrendatário
        arrendatario_crr = crr.arrendatarios.first() if crr.arrendatarios.exists() else None
        arrendatario = arrendatario_crr.arrendatario if arrendatario_crr else None
        nome_arrendatario = arrendatario.nome_arrendatario if arrendatario else ''
        
        # Preenche células
        row_cells[0].text = str(placa).upper()
        row_cells[1].text = f"{str(marca)}/{str(modelo)}".upper()
        row_cells[2].text = str(responsavel).upper()
        row_cells[3].text = str(nome_arrendatario).upper()

    # Gera resposta
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    filename = f'edital_{numero_edital}_{timezone.now().strftime("%Y%m%d_%H%M%S")}.docx'
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    return response