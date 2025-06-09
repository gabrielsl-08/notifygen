# utils.py
# Mantenha todos os imports que você já tem
from docx import Document
from docx.shared import Inches, Pt # Importe Pt se estiver usando
from django.http import HttpResponse
from io import BytesIO
from django.utils.timezone import now
from datetime import datetime
import locale
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn # Importe qn se estiver usando
import os # Importe os para lidar com caminhos
from django.conf import settings # Importe settings para acessar MEDIA_ROOT
from django.db.models import QuerySet # Importe QuerySet para type hinting

# Importe suas models necessárias. Ajuste 'notificacao.models' para o nome do seu app
from notificacao.models import NumeroEdital, Crr # Importe Crr para type hinting e prefetch_related

# Defina o caminho para o template DOCX usando settings.MEDIA_ROOT
TEMPLATE_DOCX_PATH = os.path.join(settings.MEDIA_ROOT, 'modelo_edital.docx')
# Verifique se este caminho está correto e se o arquivo existe.

# Sua função para obter o próximo número (sem atomicidade)
def obter_proximo_numero_edital():
    # Nota: Esta função tem uma condição de corrida potencial se chamada
    # por múltiplos processos ao mesmo tempo sem um lock.
    obj, _ = NumeroEdital.objects.get_or_create(id=1)
    return obj

# Sua função para gerar o DOCX (inclui a obtenção e incremento do número)
def gerar_edital_docx(crrs_queryset: QuerySet[Crr]): # Adicione type hinting
    """
    Gera um documento DOCX para um queryset de CRRs usando um template DOCX.
    Substitui placeholders, adiciona uma tabela, obtém e incrementa o contador.
    """
    # Carrega o template existente
    try:
        doc = Document(TEMPLATE_DOCX_PATH)
    except Exception as e:
        # Loga o erro e re-lança uma exceção mais amigável
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f"Erro ao carregar o template DOCX em {TEMPLATE_DOCX_PATH}: {e}")
        # Re-lança a exceção para ser tratada na view
        raise FileNotFoundError(f"Template DOCX não encontrado ou inválido em {TEMPLATE_DOCX_PATH}. Verifique settings.MEDIA_ROOT e o caminho.") from e


    # Obtém o objeto contador e o número atual
    numero_edital_obj = obter_proximo_numero_edital()
    numero_edital = str(numero_edital_obj.numero) # O número para o edital atual

    # Configura o locale para formatação da data
    try:
        locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8') # Para sistemas baseados em Unix
    except locale.Error:
        try:
             locale.setlocale(locale.LC_TIME, 'pt_BR') # Para Windows
        except locale.Error:
             pass # Continua sem locale específico se ambos falharem

    # Formata a data atual
    data_formatada = now().strftime('%d DE %B DE %Y').upper()


    # Substitui os placeholders {{DATA_ATUAL}} e {{NUMERO_EDITAL}} nos parágrafos
    # Melhorando a lógica de substituição para tentar preservar a formatação
    for paragraph in doc.paragraphs:
         # Captura o texto original do parágrafo antes de qualquer substituição
         original_text = "".join([run.text for run in paragraph.runs])
         replaced_text = original_text.replace("{{DATA_ATUAL}}", data_formatada).replace("{{NUMERO_EDITAL}}", numero_edital)

         if original_text != replaced_text:
              # Se houve substituição, limpa os runs existentes e adiciona um novo com o texto substituído
              # Tentativa de preservar alguma formatação do primeiro run
              if paragraph.runs:
                  first_run = paragraph.runs[0]
                  font_name = first_run.font.name
                  font_size = first_run.font.size
                  bold = first_run.bold
                  italic = first_run.italic
                  underline = first_run.underline
                  # Adicione outras propriedades de formatação conforme necessário

                  # Limpa os runs existentes
                  while len(paragraph.runs) > 0:
                      p = paragraph._element
                      p.remove(p.r_lst[-1]) # Remove o último run element

                  # Adiciona um novo run com o texto substituído e tenta re-aplicar a formatação
                  new_run = paragraph.add_run(replaced_text)
                  new_run.font.name = font_name
                  new_run.font.size = font_size
                  new_run.bold = bold
                  new_run.italic = italic
                  new_run.underline = underline
                  # Tenta aplicar a formatação qn se necessário (baseado no seu código original)
                  try:
                      if new_run._element.rPr is None:
                          new_run._element.add_rPr()
                      if new_run._element.rPr.rFonts is None:
                          new_run._element.rPr.add_rFonts()
                      new_run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name or 'Verdana') # Use font_name ou fallback
                  except Exception:
                       pass # Ignora se qn falhar ou rPr/rFonts não existirem


         # Aplica alinhamento central se o parágrafo original continha algum placeholder
         if "{{DATA_ATUAL}}" in original_text or "{{NUMERO_EDITAL}}" in original_text:
              paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    # Adiciona a tabela no final do documento
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid' # Deixa as bordas visíveis

    # Define cabeçalho da tabela
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'PLACA/CHASSI'
    hdr_cells[1].text = 'MARCA/MODELO'
    hdr_cells[2].text = 'RESPONSÁVEL'
    hdr_cells[3].text = 'AGENTE FINACEIRO/ARRENDATÁRIO'

    # Popula a tabela com dados dos CRRs
    # É CRUCIAL que crrs_queryset tenha os related objects prefetched na view
    # para evitar N+1 queries aqui.
    for crr in crrs_queryset:
        row_cells = table.add_row().cells

        # Acessa objetos relacionados. Use getattr para segurança caso a relação seja None.
        # Baseado na sua lógica original: crr.veiculo.first(), crr.notificacao, crr.arrendatarios.first().arrendatario
        # Isso sugere que 'veiculo' é um related_name de uma FK em Veiculo, 'notificacao' é FK/OTO em Crr,
        # e 'arrendatarios' é ManyToMany via intermediária.

        # Acessando Veiculo: crr.veiculo.first() implica que 'veiculo' é um RelatedManager
        veiculo_obj = getattr(crr, 'veiculo', None)
        if veiculo_obj and isinstance(veiculo_obj, QuerySet): # Verifica se é um related manager
             veiculo_obj = veiculo_obj.first() # Pega o primeiro objeto relacionado

        placa = getattr(veiculo_obj, 'placa', '') if veiculo_obj else ''
        marca = getattr(veiculo_obj, 'marca', '') if veiculo_obj else ''
        modelo = getattr(veiculo_obj, 'modelo', '') if veiculo_obj else ''

        # Acessando Notificacao: crr.notificacao implica que 'notificacao' é o objeto FK/OTO
        notificacao_obj = getattr(crr, 'notificacao', None)
        destinatario = getattr(notificacao_obj, 'destinatario', '') if notificacao_obj else ''

        # Acessando Arrendatario: crr.arrendatarios.first().arrendatario implica ManyToMany via intermediária
        arrendatarios_manager = getattr(crr, 'arrendatarios', None) # Pega o related manager
        nome_arrendatario = ''
        if arrendatarios_manager and isinstance(arrendatarios_manager, QuerySet): # Verifica se é um related manager
             intermediate_obj = arrendatarios_manager.first() # Pega o primeiro objeto intermediário
             if intermediate_obj:
                  final_arrendatario_obj = getattr(intermediate_obj, 'arrendatario', None) # Pega o objeto Arrendatario final
                  if final_arrendatario_obj:
                       nome_arrendatario = getattr(final_arrendatario_obj, 'nome_arrendatario', '')


        row_cells[0].text = str(placa).upper() # Garante que é string antes de upper
        row_cells[1].text = f"{str(marca)}/{str(modelo)}".upper() # Garante que são strings
        row_cells[2].text = str(destinatario).upper() # Garante que é string
        row_cells[3].text = str(nome_arrendatario).upper() # Garante que é string


    # Salvar o arquivo em memória
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Criar resposta para download
    # O nome do arquivo inclui o número do edital gerado
    filename = f'edital_{numero_edital}_{timezone.now().strftime("%Y%m%d_%H%M%S")}.docx'

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'

    # *** CHAMA O INCREMENTO AQUI, COMO NA SUA FUNÇÃO ORIGINAL ***
    # Isto incrementa o contador NumeroEdital para o PRÓXIMO uso.
    # Nota: Esta chamada está DENTRO da função de geração do DOCX,
    # replicando o comportamento da sua função original.
    # Se esta função falhar após obter o número mas antes de incrementar,
    # o número será "perdido". Se múltiplos usuários chamarem ao mesmo tempo,
    # pode haver números duplicados.
    try:
        numero_edital_obj.incrementar()
    except Exception as e:
         # Loga o erro, mas não impede o download do arquivo gerado
         logger.error(f"Erro ao incrementar o contador NumeroEdital: {e}")
         # Considere adicionar um sistema de alerta aqui se isso for crítico.

    return response
